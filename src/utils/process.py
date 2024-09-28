# Read and process the translated text file, then convert it to a PDF file

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Pt

# Processes a text file by removing specific text and formatting
def process_file(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Remove the specified text
    content = content.replace("Here is the English translation of the Hebrew text, preserving all formatting:", "")
    content = content.replace("Here is the translation from Hebrew to English, preserving the formatting:", "")
    content = content.replace("Here is the English translation of the Hebrew text, preserving the formatting:", "")
    content = content.replace("Here is the translation from Hebrew to English, preserving all formatting:", "")
    content = content.replace("Here is the English translation of the Hebrew text, with formatting preserved:", "")
    content = content.replace("Here is the English translation, preserving the formatting:", "")
    content = content.replace("Here is the English translation, preserving all formatting:", "")
    
    # Remove "--- New Page ---"
    content = content.replace("--- New Page ---", "")
    
    # Remove extra blank lines
    content = '\n'.join(line for line in content.splitlines() if line.strip())
    
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write(content)

# Converts a text file to a PDF file with specific formatting
def text_to_pdf(input_file, output_file):
    # Read the input file
    with open(input_file, 'r', encoding='utf-8') as file:
        content = file.read().split('\n')

    # Create a PDF document
    doc = SimpleDocTemplate(output_file, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    page_number = 0
    first_page_found = False

    # Process each line
    for line in content:
        line = line.strip()
        if line.isdigit():
            if not first_page_found and line == '1':
                first_page_found = True
            if first_page_found:
                if page_number > 0:
                    story.append(PageBreak())
                page_number += 1
                story.append(Paragraph(f"Page {page_number}", styles['Normal']))
        elif line:
            if first_page_found:
                story.append(Paragraph(line, styles['Normal']))
    
    # Build the PDF
    doc.build(story)

# Converts a text file to a Word document, preserving all formatting
def text_to_word(input_file, output_file):
    # Read the input file
    with open(input_file, 'r', encoding='utf-8') as file:
        content = file.read().split('\n')

    # Create a Word document
    doc = Document()
    page_number = 0
    first_page_found = False
    current_page_content = []

    # Process each line
    for line in content:
        # Preserve leading and trailing spaces
        line = line.rstrip()
        if line.isdigit():
            if not first_page_found and line == '1':
                first_page_found = True
            if first_page_found:
                if page_number > 0:
                    # Check if the current page has content before adding a page break
                    if current_page_content:
                        doc.add_page_break()
                        current_page_content = []
                page_number += 1
                paragraph = doc.add_paragraph(f"Page {page_number}")
                paragraph.style.font.size = Pt(12)
        else:
            if first_page_found:
                # Add paragraph with preserved formatting
                paragraph = doc.add_paragraph(line)
                paragraph.style.font.size = Pt(12)
                current_page_content.append(line)
    
    # Remove any trailing blank pages
    if not current_page_content:
        # If the last page is blank, remove the last page break
        doc.paragraphs[-1].clear()

    # Save the Word document
    doc.save(output_file)
    
# Removes all fully blank pages from a Word document and saves the result to a new file.
def remove_blank_pages(input_file, output_file):
    # Load the input Word document
    doc = Document(input_file)
    new_doc = Document()
    
    # Track if the current page is blank
    is_blank_page = True
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            is_blank_page = False
            new_doc.add_paragraph(paragraph.text)
        elif not is_blank_page:
            new_doc.add_paragraph(paragraph.text)
    
    # Save the new document without blank pages
    new_doc.save(output_file)

def remove_spaces_between_lines(input_file, output_file):
    # Load the input Word document
    doc = Document(input_file)
    new_doc = Document()
    
    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        # Strip leading and trailing spaces from the paragraph text
        stripped_text = paragraph.text.strip()
        if stripped_text:
            # Add the stripped text to the new document
            new_doc.add_paragraph(stripped_text)
    
    # Save the new document without spaces between lines
    new_doc.save(output_file)

def add_new_page_on_marker(input_file, output_file):
    # Load the input Word document
    doc = Document(input_file)
    new_doc = Document()
    
    # Regular expression to match "Page x" where x is an integer
    import re
    page_marker_pattern = re.compile(r'^Page \d+$')
    
    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        # Check if the paragraph matches the page marker pattern
        if page_marker_pattern.match(paragraph.text.strip()):
            # Add a page break before adding the paragraph to the new document
            new_doc.add_page_break()
        # Add the paragraph to the new document
        new_doc.add_paragraph(paragraph.text)
    
    # Save the new document with page breaks
    new_doc.save(output_file)



if __name__ == "__main__":
    input_file = "src/data/final1.docx"
    output_file = "src/data/final2.docx"
    # process_file(input_file, output_file)
    # text_to_word(input_file, output_file)
    # remove_blank_pages(input_file, output_file)
    remove_spaces_between_lines(input_file, output_file)
    print(f"Processed file saved as {output_file}")